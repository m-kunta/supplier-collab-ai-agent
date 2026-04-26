import re
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX

def markdown_to_docx(md_text, output_path):
    doc = Document()
    
    in_table = False
    table_data = []
    
    def process_table():
        if not table_data: return
        # First row is header, second is separator
        if len(table_data) > 1 and all(c.replace('-', '').strip() == '' for c in table_data[1]):
            table_data.pop(1)
            
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            row_text_lower = " ".join(row_data).lower()
            
            bg_color = None
            if i > 0: # Not header
                if 'red' in row_text_lower and 'yellow' not in row_text_lower and 'green' not in row_text_lower:
                    bg_color = RGBColor(255, 230, 230) # Light red
                elif 'yellow' in row_text_lower and 'green' not in row_text_lower:
                    bg_color = RGBColor(255, 255, 200) # Light yellow
                elif 'green' in row_text_lower:
                    bg_color = RGBColor(230, 255, 230) # Light green
                    
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    cell = row.cells[j]
                    p = cell.paragraphs[0]
                    # Bold header
                    if i == 0:
                        run = p.add_run(cell_text.strip())
                        run.bold = True
                    else:
                        p.add_run(cell_text.strip())
                    
                    if bg_color:
                        from docx.oxml.shared import OxmlElement
                        from docx.oxml.ns import qn
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), '%02x%02x%02x' % (bg_color[0], bg_color[1], bg_color[2]))
                        cell._tc.get_or_add_tcPr().append(shd)
                        
        table_data.clear()

    for line in md_text.split('\n'):
        stripped = line.strip()
        
        if stripped.startswith('|'):
            in_table = True
            cols = [c.strip() for c in stripped.split('|')[1:-1]]
            table_data.append(cols)
            continue
            
        if in_table:
            process_table()
            in_table = False
            
        if not stripped:
            continue
            
        if stripped.startswith('---'):
            continue  # Skip front-matter separators
            
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)

    if in_table:
        process_table()
        
    doc.save(output_path)
    print(f"Saved to {output_path}")

md_content = Path("output/V1001_2026-04-18.md").read_text()
import re
md_content = re.sub(r'^---[\s\S]*?---\n', '', md_content)
markdown_to_docx(md_content, "output/test.docx")
