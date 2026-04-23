import os
import json
from pathlib import Path
from unittest.mock import MagicMock
import pytest

from src.output_renderer import render_markdown, render_docx, write_output

# Local mocks so we don't need a full agent setup
class MockProvider:
    provider = "mock_provider"
    model = "mock_model"

class MockContext:
    def __init__(self):
        self.vendor_id = "V123"
        self.meeting_date = "2026-04-18"
        self.persona_emphasis = "both"
        self.provider = MockProvider()
        self.pipeline_notes = []
        self.validation_report = {
            "overall_status": "passed",
            "error_count": 0,
            "warning_count": 0,
            "manifest": {"status": "passed", "errors": [], "warnings": []},
            "datasets": {"vendor_master": {"status": "passed", "errors": [], "warnings": []}},
        }
        self.config = {"output": {"footer_text": "Mock Footer"}}
        self.briefing_text = """
# Heading 1
## Heading 2

Some **bold text** here.

| Header 1 | Header 2 | Risk |
|---|---|---|
| Value 1 | Value 2 | **Red** |
| Value 3 | Value 4 | Green |

- Bullet 1
- Bullet 2
"""

@pytest.fixture
def mock_ctx():
    return MockContext()

def test_render_markdown(mock_ctx):
    output = render_markdown(mock_ctx)
    assert "vendor_id: V123" in output
    assert "Mock Footer" in output
    assert "# Heading 1" in output

def test_render_docx_success(mock_ctx, tmp_path):
    output_path = tmp_path / "test.docx"
    result = render_docx(mock_ctx, output_path)
    
    assert result == output_path
    assert output_path.exists()
    assert output_path.stat().st_size > 0

    from docx import Document
    doc = Document(output_path)
    
    # Check headings and paragraphs
    paragraphs = [p.text for p in doc.paragraphs]
    assert "Heading 1" in paragraphs
    assert "Heading 2" in paragraphs
    assert "Some bold text here." in paragraphs
    assert any("Mock Footer" in p for p in paragraphs)
    
    # Check tables
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 3
    assert table.cell(0, 0).text == "Header 1"
    assert table.cell(1, 2).text == "Red"

def test_write_output_md(mock_ctx, tmp_path):
    result = write_output(mock_ctx, tmp_path, "md")
    assert "md_path" in result
    assert "validation_report_path" in result
    assert "docx_path" not in result
    assert result["md_path"].exists()
    assert result["validation_report_path"].exists()

    report = json.loads(result["validation_report_path"].read_text(encoding="utf-8"))
    assert report["overall_status"] == "passed"

def test_write_output_both(mock_ctx, tmp_path):
    result = write_output(mock_ctx, tmp_path, "both")
    assert "md_path" in result
    assert "docx_path" in result
    assert "validation_report_path" in result
    assert result["md_path"].exists()
    assert result["docx_path"].exists()
    assert result["validation_report_path"].exists()

def test_render_docx_empty_text_raises():
    ctx = MockContext()
    ctx.briefing_text = ""
    with pytest.raises(ValueError, match="is empty"):
        render_docx(ctx, Path("dummy.docx"))
